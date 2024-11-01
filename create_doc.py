from qgis.core import (
    QgsApplication,
    QgsProject,
    QgsVectorLayer,
    QgsCoordinateReferenceSystem,
    QgsCoordinateTransform,
    QgsSymbol,
    QgsSingleSymbolRenderer,
    edit,
    QgsLayout,
    QgsLayoutItemMap,
    QgsLayoutExporter,
    QgsLayoutSize,
    QgsLayoutPoint
)
from PyQt5.QtWidgets import QApplication, QFileDialog
from PyQt5.QtCore import QRectF
from docx import Document
from docx.shared import Inches
import os
import sys
import datetime

# Function to get the CSV directory path
def get_csv_dir():
    app = QApplication(sys.argv)
    csv_dir = QFileDialog.getExistingDirectory(None, "Select CSV Directory")
    app.exit()
    return csv_dir

# Get the CSV directory path
csv_dir = get_csv_dir()

# Define paths
# csv_dir = "C:/Users/PC/Desktop/zeghbaba"
image_output_folder = "C:/Users/PC/Desktop/screenshots"

# Your existing code using csv_dir
print(f"Selected CSV directory: {csv_dir}")

# Paths to the .qml style files
style_qml_1 = r'C:\Users\PC\Desktop\stage\projet\Styles\style_GSM.qml'
style_qml_2 = r'C:\Users\PC\Desktop\stage\projet\Styles\style_UMTS.qml'

# Identifiers for the styles
identifiers_1 = ['GSM', 'G1800', 'G900']
identifiers_2 = ['UMTS', 'U2100', 'U900']

# Ensure the directories exist
os.makedirs(image_output_folder, exist_ok=True)

# Initialize QGIS Application
QgsApplication.setPrefixPath("E:/Program Files/QGIS 3.38.0", True)
qgs = QgsApplication([], False)
qgs.initQgis()

# Function to open a file dialog to select the QGIS project file
def select_qgis_project_file():
    options = QFileDialog.Options()
    options |= QFileDialog.DontUseNativeDialog
    file_dialog = QFileDialog()
    file_dialog.setAcceptMode(QFileDialog.AcceptOpen)
    file_dialog.setNameFilter("QGIS Projects (*.qgz)")
    file_dialog.setFileMode(QFileDialog.ExistingFiles)
    file_dialog.setViewMode(QFileDialog.List)
    project_path, _ = file_dialog.getOpenFileName(None, "Select QGIS Project File", "", "QGIS Projects (*.qgz);;All Files (*)", options=options)
    return project_path

# Open the QGIS project
qgis_project_path = select_qgis_project_file()
if not qgis_project_path:
    print("No QGIS project file selected. Exiting.")
    sys.exit(1)

project = QgsProject.instance()
if not project.read(qgis_project_path):
    print(f"Failed to read the project: {qgis_project_path}")
    sys.exit(1)

def add_csv_layer(csv_path):
    uri = f"file:///{csv_path}?delimiter=,&xField=Longitude&yField=Latitude"
    layer_name = os.path.basename(csv_path).replace('.csv', '')
    layer = QgsVectorLayer(uri, layer_name, "delimitedtext")
    
    if not layer.isValid():
        print(f"Failed to load {csv_path}")
        return False

    # Set the layer's CRS to WGS 84 (EPSG:4326)
    layer.setCrs(QgsCoordinateReferenceSystem("EPSG:4326"))
    
    # Add the layer to the project
    project.addMapLayer(layer)
    return True

def add_shapefile_layer(shapefile_path, layer_name):
    layer = QgsVectorLayer(shapefile_path, layer_name, 'ogr')
    
    if not layer.isValid():
        print(f"Failed to load {shapefile_path}")
        return False

    # Add the layer to the project
    project.addMapLayer(layer)
    return True

def transform_layer_crs(layer, target_crs_authid):
    source_crs = QgsCoordinateReferenceSystem(layer.crs().authid())
    target_crs = QgsCoordinateReferenceSystem(target_crs_authid)
    with edit(layer):
        for feat in layer.getFeatures():
            geom = feat.geometry()
            geom.transform(QgsCoordinateTransform(source_crs, target_crs, QgsProject.instance()))
            feat.setGeometry(geom)
            layer.updateFeature(feat)
        layer.setCrs(target_crs)

# Iterate through CSV files in the directory and add them as layers
for csv_file in os.listdir(csv_dir):
    if csv_file.endswith(".csv"):
        csv_path = os.path.join(csv_dir, csv_file)
        if add_csv_layer(csv_path):
            print(f"Successfully added {csv_path}")
        else:
            print(f"Failed to add {csv_path}")

# Apply .qml styles based on identifiers
layers = QgsProject.instance().mapLayers().values()
for layer in layers:
    if isinstance(layer, QgsVectorLayer):
        layer_name = layer.name()
        if any(identifier in layer_name for identifier in identifiers_1):
            if layer.loadNamedStyle(style_qml_1):
                print(f"Style (GSM/G1800/G900) applied to layer: {layer_name}")
            else:
                print(f"Failed to apply GSM/G1800/G900 style to layer: {layer_name}")
        elif any(identifier in layer_name for identifier in identifiers_2):
            if layer.loadNamedStyle(style_qml_2):
                print(f"Style (UMTS/U2100/U900) applied to layer: {layer_name}")
            else:
                print(f"Failed to apply UMTS/U2100/U900 style to layer: {layer_name}")

# Create a new Word document
doc = Document()
doc.add_heading(f'Map Layers from Project: {os.path.basename(qgis_project_path)}', level=1)

# Function to export a layout to an image
def export_layout_to_image(layout, output_image_path):
    try:
        exporter = QgsLayoutExporter(layout)
        result = exporter.exportToImage(output_image_path, QgsLayoutExporter.ImageExportSettings())

        if result == QgsLayoutExporter.Success:
            print(f"Successfully exported layout to {output_image_path}")
            return True
        else:
            print(f"Failed to export layout: {result}")
            return False
    except Exception as e:
        print(f"Failed to export layout: {e}")
        return False

# Separate layers into TUNISIA layers and other layers
tunisian_layers = [layer for layer in project.mapLayers().values() if "TUNISIA" in layer.name()]
other_layers = [layer for layer in project.mapLayers().values() if "TUNISIA" not in layer.name()]

# Iterate through all the layers in the project
for layer in other_layers:
    # Create a new layout and map item for each layer
    layout = QgsLayout(project)
    layout.initializeDefaults()
    
    # Add a map item to the layout
    layout_item_map = QgsLayoutItemMap(layout)
    layout_item_map.setRect(QRectF(20, 20, 200, 200))
    layout.addLayoutItem(layout_item_map)

    # Ensure TUNISIA layers are always visible
    layout_item_map.setLayers(tunisian_layers + [layer])

    # Set the extent to the layer's extent
    layer_extent = layer.extent()
    layout_item_map.setExtent(layer_extent)

    # Set size and position
    layout_item_map.attemptMove(QgsLayoutPoint(20, 20))
    layout_item_map.attemptResize(QgsLayoutSize(200, 200))

    # Export the layout to an image
    screenshot_path = os.path.join(image_output_folder, f"{layer.name()}.png")
    if export_layout_to_image(layout, screenshot_path):
        # Add the screenshot to the Word document
        doc.add_heading(layer.name(), level=2)
        doc.add_picture(screenshot_path, width=Inches(6))
    else:
        print(f"Skipping {layer.name()} due to export failure.")

# Automatically save the project
if not qgis_project_path.endswith('.qgz'):
    qgis_project_path += '.qgz'
QgsProject.instance().write(qgis_project_path)
print(f'Project saved to {qgis_project_path}.')

# Generate the output Word document path based on the project save path
project_save_dir = os.path.dirname(qgis_project_path)
current_time = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
output_word_path = os.path.join(project_save_dir, f"test_{current_time}.docx")

# Save the Word document
doc.save(output_word_path)
print(f"Document saved to {output_word_path}")

